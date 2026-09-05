"""
POST /tools/search-docs  (contract §2.6)

Local vector search over a small corpus of SOPs/manuals in docs_corpus/,
plus anything supported dropped into FILES_DIR (the same directory
generate_file writes to and Person B serves at /files/) — so any document a
user places there becomes searchable, not just the curated sample SOPs.
Uses ChromaDB with its default (local, on-disk, no separate server process)
embedding + storage. Everything here is local-only — no external API calls,
consistent with the air-gapped requirement.
"""
import hashlib
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional, Tuple

import chromadb
from chromadb.utils.embedding_functions.onnx_mini_lm_l6_v2 import ONNXMiniLM_L6_V2
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from .config import REPO_ROOT, FILES_DIR

# Force the CPU execution provider for the local embedding model. Chroma's
# default provider selection tries CoreML first on Apple Silicon, which
# fails hard (onnxruntime CoreML execution provider error) in some sandboxed/
# restricted macOS environments — an environment-compatibility fix, not a
# contract change (the search-docs request/response shape is untouched).
_EMBEDDING_FUNCTION = ONNXMiniLM_L6_V2(preferred_providers=["CPUExecutionProvider"])

CORPUS_DIR = REPO_ROOT / "tools" / "docs_corpus"
CHROMA_DIR = REPO_ROOT / "tools" / "chroma_db"
COLLECTION_NAME = "mrpl_docs"

# Chat-scoped Knowledge Base. A SECOND collection inside the SAME ChromaDB
# (same PersistentClient, same on-disk store, same embedding function, same
# _chunk_text) — not a second vector database. Kept separate from
# COLLECTION_NAME purely so the directory-corpus auto-sync (_sync_corpus,
# which deletes chunks whose `source` is no longer a file on disk) can never
# touch documents that were ingested from a chat upload. Every chunk here
# carries a `chat_id` in its metadata and every query is filtered by it, so
# a document uploaded in one chat is never retrievable from another.
CHAT_COLLECTION_NAME = "chat_kb"

# Every directory that may contain searchable source documents. FILES_DIR is
# intentionally included: it's the same directory Person C's own
# generate_file writes to and Person B serves at /files/, so anything a user
# drops there (in a supported format) becomes searchable too.
SOURCE_DIRS = (CORPUS_DIR, FILES_DIR)


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


# Text extraction is only defined for these extensions (matched
# case-insensitively). Adding a new supported format means adding one entry
# here — nothing else in the ingestion pipeline needs to change.
_TEXT_EXTRACTORS = {
    ".txt": _read_txt,
    ".md": _read_txt,
    ".docx": _read_docx,
}

_client = None
_collection = None
_chat_collection = None


def _chunk_text(text: str, chunk_size: int = 800, overlap: int = 100):
    chunks = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + chunk_size, n)
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return [c.strip() for c in chunks if c.strip()]


def _discover_files() -> list[Path]:
    """
    Every supported-format file across all SOURCE_DIRS, skipping
    directories, unsupported extensions, hidden files, and Microsoft
    Office's transient "~$..." lock files (created while a document is open
    for editing elsewhere — these are short owner-info stubs, not real
    document content, and are not even valid docx/zip archives).
    """
    seen: set[Path] = set()
    files: list[Path] = []
    for source_dir in SOURCE_DIRS:
        source_dir.mkdir(parents=True, exist_ok=True)
        resolved_dir = source_dir.resolve()
        print(
            f"[docsearch] scanning source dir: {resolved_dir} "
            f"(supported extensions: {sorted(_TEXT_EXTRACTORS)})"
        )
        for path in sorted(source_dir.iterdir()):
            if not path.is_file():
                continue
            if path.name.startswith("~$") or path.name.startswith("."):
                print(f"[docsearch]   skip (lock/hidden file): {path.name!r}")
                continue
            if path.suffix.lower() not in _TEXT_EXTRACTORS:
                print(f"[docsearch]   skip (unsupported extension {path.suffix!r}): {path.name!r}")
                continue
            resolved = path.resolve()
            if resolved in seen:
                print(f"[docsearch]   skip (duplicate path across source dirs): {path.name!r}")
                continue  # e.g. CORPUS_DIR and FILES_DIR happen to coincide
            seen.add(resolved)
            print(f"[docsearch]   discovered: {path.name!r} ({resolved})")
            files.append(path)
    print(f"[docsearch] discovery complete: {len(files)} candidate file(s) across {len(SOURCE_DIRS)} source dir(s)")
    return files


def _extract_text(path: Path) -> Optional[str]:
    """
    Returns the extracted plain text for `path`, or None if it couldn't be
    parsed (corrupted file, truncated archive, etc). The caller skips such
    files with a warning instead of failing the whole ingestion pass — one
    bad document must never take down search for every other document.
    """
    extractor = _TEXT_EXTRACTORS.get(path.suffix.lower())
    if extractor is None:
        return None
    try:
        text = extractor(path)
        print(f"[docsearch]   parsed {path.name!r}: {len(text)} chars extracted")
        return text
    except (PackageNotFoundError, OSError, ValueError) as e:
        print(f"[docsearch]   PARSE FAILED for {path.name!r} ({type(e).__name__}): {e}")
        return None
    except Exception as e:  # noqa: BLE001 - one bad file must never break ingestion
        print(f"[docsearch]   PARSE FAILED for {path.name!r} ({type(e).__name__}): {e}")
        return None


def _existing_sources(collection) -> dict:
    """
    Maps source filename -> the content_hash it was last indexed with, from
    whatever's currently in the collection. Used to detect new, unchanged,
    modified, and deleted documents on every sync pass.
    """
    try:
        got = collection.get(include=["metadatas"])
    except Exception:
        return {}
    sources: dict = {}
    for meta in got.get("metadatas") or []:
        if not meta:
            continue
        source = meta.get("source")
        if source and source not in sources:
            sources[source] = meta.get("content_hash", "")
    return sources


def _sync_corpus(collection):
    """
    Re-syncs the collection against the current state of SOURCE_DIRS. Safe
    to call on every request: unchanged files are a no-op, new files are
    ingested, modified files (detected via a content hash, not just
    filename) are re-ingested — their stale chunks are removed before the
    fresh ones are added — and files removed from disk have their leftover
    chunks removed too, so the index never accumulates stale or orphaned
    data. Idempotent, and independent of process cwd (SOURCE_DIRS are
    resolved from REPO_ROOT-anchored config, see app/config.py).
    """
    print(f"[docsearch] --- sync pass starting (collection count before: {collection.count()}) ---")
    discovered = _discover_files()
    discovered_names = {path.name for path in discovered}
    existing_hashes = _existing_sources(collection)
    print(f"[docsearch] sources already in index: {sorted(existing_hashes)}")

    # Drop chunks for any source that's indexed but no longer on disk.
    for source in existing_hashes:
        if source not in discovered_names:
            print(f"[docsearch] source removed from disk, deleting its chunks: {source!r}")
            try:
                collection.delete(where={"source": source})
            except Exception as e:  # noqa: BLE001
                print(f"[docsearch] failed to remove stale source {source!r}: {e}")

    ids, docs, metadatas = [], [], []
    for path in discovered:
        text = _extract_text(path)
        if not text or not text.strip():
            print(f"[docsearch]   no usable text extracted from {path.name!r} -> skipping")
            continue

        content_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]
        if existing_hashes.get(path.name) == content_hash:
            print(f"[docsearch]   {path.name!r} unchanged (hash={content_hash}) -> no-op")
            continue  # already indexed with this exact content — no-op

        if path.name in existing_hashes:
            # Content changed since it was last indexed — clear the stale
            # chunks first so old and new text never coexist for one source.
            print(
                f"[docsearch]   {path.name!r} content changed "
                f"(old hash={existing_hashes[path.name]!r}, new hash={content_hash}) -> reindexing"
            )
            try:
                collection.delete(where={"source": path.name})
            except Exception as e:  # noqa: BLE001
                print(f"[docsearch] failed to clear stale chunks for {path.name!r}: {e}")
        else:
            print(f"[docsearch]   {path.name!r} is new (hash={content_hash}) -> indexing")

        file_chunks = _chunk_text(text)
        print(f"[docsearch]   {path.name!r} chunked into {len(file_chunks)} chunk(s)")
        for i, chunk in enumerate(file_chunks):
            ids.append(f"{path.name}::{content_hash}::{i}")
            docs.append(chunk)
            metadatas.append({"source": path.name, "content_hash": content_hash})

    if ids:
        print(f"[docsearch] adding {len(ids)} new chunk embedding(s) to the index across "
              f"{len({m['source'] for m in metadatas})} source(s)")
        collection.add(ids=ids, documents=docs, metadatas=metadatas)
    else:
        print("[docsearch] no new/changed chunks to add this pass")
    print(f"[docsearch] --- sync pass complete (collection count after: {collection.count()}) ---")


def _get_client():
    """The single shared PersistentClient for both collections (corpus + chat KB)."""
    global _client
    if _client is None:
        CHROMA_DIR.mkdir(parents=True, exist_ok=True)
        print(f"[docsearch] opening persistent Chroma index at: {CHROMA_DIR.resolve()}")
        _client = chromadb.PersistentClient(path=str(CHROMA_DIR))
    return _client


def get_collection():
    global _collection
    if _collection is None:
        _collection = _get_client().get_or_create_collection(
            name=COLLECTION_NAME, embedding_function=_EMBEDDING_FUNCTION
        )
        print(f"[docsearch] collection {COLLECTION_NAME!r} loaded, "
              f"{_collection.count()} chunk(s) already persisted on disk")

    # Re-sync on every call (not just the first) so documents added,
    # modified, or removed while the service is already running are picked
    # up without requiring a restart.
    _sync_corpus(_collection)
    return _collection


def search_docs(query: str, top_k: int = 3) -> list[dict]:
    collection = get_collection()
    count = collection.count()
    if count == 0:
        return []

    top_k = max(1, min(top_k, count))
    result = collection.query(query_texts=[query], n_results=top_k)

    out = []
    docs = result.get("documents", [[]])[0]
    metas = result.get("metadatas", [[]])[0]
    dists = result.get("distances", [[]])[0]
    for text, meta, dist in zip(docs, metas, dists):
        # Chroma returns a distance (lower = closer); convert to a
        # 0-1 "similarity-ish" score for the contract's `score` field.
        score = 1.0 / (1.0 + dist) if dist is not None else 0.0
        out.append({
            "text": text,
            "source": (meta or {}).get("source", "unknown"),
            "score": round(float(score), 4),
        })
    return out


# ---------------------------------------------------------------------------
# Chat-scoped Knowledge Base (the `chat_kb` collection)
# ---------------------------------------------------------------------------

def _now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def get_chat_collection():
    global _chat_collection
    if _chat_collection is None:
        _chat_collection = _get_client().get_or_create_collection(
            name=CHAT_COLLECTION_NAME, embedding_function=_EMBEDDING_FUNCTION
        )
        print(
            f"[docsearch] chat KB collection {CHAT_COLLECTION_NAME!r} loaded, "
            f"{_chat_collection.count()} chunk(s) already persisted on disk"
        )
    return _chat_collection


def ingest_chat_document(
    chat_id: str,
    filename: str,
    pages: List[Tuple[int, str]],
    document_id: Optional[str] = None,
    chat_title: Optional[str] = None,
) -> dict:
    """
    Chunk + embed + persist an uploaded document into the chat_kb collection,
    every chunk tagged with the owning chat_id (mandatory isolation key),
    document_id, filename and page number.

    `pages` is [(page_number, page_text), ...] as returned by
    pdf_ingest.extract_pdf_pages — chunking is done per page so page numbers
    survive into each chunk's metadata for citations.

    Re-uploading the same filename into the same chat replaces that
    document's existing chunks (mirrors the corpus sync's content-change
    reindex behavior) rather than duplicating them.

    Returns { document_id, filename, chat_id, chunks, status }.
    Raises ValueError if no non-empty text chunks could be produced.
    """
    if not chat_id or not chat_id.strip():
        raise ValueError("chat_id is required for chat-scoped ingestion")
    chat_id = chat_id.strip()
    document_id = document_id or str(uuid.uuid4())

    collection = get_chat_collection()

    # Duplicate handling: drop any prior chunks for this (chat_id, filename).
    try:
        collection.delete(where={"$and": [{"chat_id": chat_id}, {"filename": filename}]})
    except Exception as e:  # noqa: BLE001
        print(f"[docsearch] chat KB: could not clear prior chunks for {filename!r} in {chat_id}: {e}")

    uploaded_at = _now_iso()
    ids: list[str] = []
    docs: list[str] = []
    metas: list[dict] = []
    chunk_index = 0
    for page_number, page_text in pages:
        for chunk in _chunk_text(page_text):
            meta = {
                "chat_id": chat_id,
                "document_id": document_id,
                "filename": filename,
                "source": filename,  # keep `source` populated too, like the corpus schema
                "chunk_index": chunk_index,
                "uploaded_at": uploaded_at,
            }
            if page_number is not None:
                meta["page"] = int(page_number)
            if chat_title:
                meta["chat_title"] = chat_title
            ids.append(f"{chat_id}::{document_id}::{chunk_index}")
            docs.append(chunk)
            metas.append(meta)
            chunk_index += 1

    if not ids:
        raise ValueError("no extractable text content in document")

    print(
        f"[docsearch] chat KB: adding {len(ids)} chunk(s) for document {filename!r} "
        f"(document_id={document_id}) in chat {chat_id}"
    )
    collection.add(ids=ids, documents=docs, metadatas=metas)
    return {
        "document_id": document_id,
        "filename": filename,
        "chat_id": chat_id,
        "chunks": len(ids),
        "status": "indexed",
    }


def search_chat_docs(query: str, chat_id: str, top_k: int = 4) -> list[dict]:
    """
    Vector search restricted to ONE chat's uploaded documents.
    The `where={"chat_id": chat_id}` filter is mandatory — a document
    uploaded in another chat can never be returned here.
    """
    if not chat_id or not chat_id.strip():
        return []
    chat_id = chat_id.strip()
    collection = get_chat_collection()
    try:
        count = collection.count()
    except Exception:  # noqa: BLE001
        count = 0
    if count == 0:
        return []

    top_k = max(1, min(top_k, count))
    result = collection.query(
        query_texts=[query],
        n_results=top_k,
        where={"chat_id": chat_id},
    )

    out = []
    docs = result.get("documents", [[]])[0]
    metas = result.get("metadatas", [[]])[0]
    dists = result.get("distances", [[]])[0]
    for text, meta, dist in zip(docs, metas, dists):
        meta = meta or {}
        score = 1.0 / (1.0 + dist) if dist is not None else 0.0
        out.append({
            "text": text,
            "source": meta.get("filename") or meta.get("source", "unknown"),
            "page": meta.get("page"),
            "document_id": meta.get("document_id"),
            "score": round(float(score), 4),
        })
    return out


def list_chat_documents(chat_id: Optional[str] = None) -> list[dict]:
    """
    Aggregates the chat_kb collection's chunk metadata into a per-document
    listing for the Admin -> Knowledge Base view. No separate table — the
    vector store's own metadata is the single source of truth.
    """
    collection = get_chat_collection()
    where = {"chat_id": chat_id.strip()} if chat_id and chat_id.strip() else None
    try:
        got = collection.get(where=where, include=["metadatas"])
    except Exception as e:  # noqa: BLE001
        print(f"[docsearch] chat KB: list_chat_documents failed: {e}")
        return []

    by_doc: dict[str, dict] = {}
    for meta in got.get("metadatas") or []:
        if not meta:
            continue
        document_id = meta.get("document_id") or meta.get("filename")
        entry = by_doc.get(document_id)
        filename = meta.get("filename") or meta.get("source") or "unknown"
        ext = Path(filename).suffix.lower().lstrip(".") or "pdf"
        if entry is None:
            by_doc[document_id] = {
                "document_id": document_id,
                "filename": filename,
                "file_type": ext,
                "chat_id": meta.get("chat_id"),
                "chat_title": meta.get("chat_title"),
                "uploaded_at": meta.get("uploaded_at"),
                "chunks": 1,
                "status": "indexed",
            }
        else:
            entry["chunks"] += 1
            if not entry.get("chat_title") and meta.get("chat_title"):
                entry["chat_title"] = meta.get("chat_title")

    return sorted(
        by_doc.values(),
        key=lambda d: (d.get("uploaded_at") or ""),
        reverse=True,
    )
